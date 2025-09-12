#!/usr/bin/env python3
"""
Document Processor Service for SCAD GenAI Data Query
Handles OCR and text extraction from various document types
"""

import os
import io
import logging
from typing import Dict, List, Optional, Union, Any
from pathlib import Path
import tempfile
import magic

# Core document processing
import pandas as pd
import pdfplumber
from docx import Document
import easyocr
from PIL import Image

# File type detection
import mimetypes

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Comprehensive document processor supporting:
    - PDFs (text extraction + OCR fallback)
    - Word documents (.docx)
    - Excel files (.xlsx, .xls)
    - CSV files
    - Images (OCR)
    """
    
    def __init__(self):
        """Initialize the document processor with OCR capabilities."""
        self.ocr_reader = None
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_excel,
            'application/vnd.ms-excel': self._process_excel_old,
            'text/csv': self._process_csv,
            'image/png': self._process_image,
            'image/jpeg': self._process_image,
            'image/jpg': self._process_image,
            'image/tiff': self._process_image,
            'image/bmp': self._process_image
        }
    
    def _init_ocr(self):
        """Lazy initialization of OCR reader."""
        if self.ocr_reader is None:
            try:
                # Initialize EasyOCR with English and Arabic support
                self.ocr_reader = easyocr.Reader(['en', 'ar'], gpu=False)
                logger.info("✅ EasyOCR initialized with English and Arabic support")
            except Exception as e:
                logger.error(f"❌ Failed to initialize OCR: {e}")
                self.ocr_reader = None
    
    def detect_file_type(self, file_path: str) -> str:
        """
        Detect file MIME type using python-magic with fallback to extension-based detection.
        
        Args:
            file_path: Path to the file
            
        Returns:
            MIME type string
        """
        try:
            mime_type = magic.from_file(file_path, mime=True)
            logger.info(f"🔍 Detected file type: {mime_type} for {file_path}")
            
            # Fix common Windows issues with python-magic
            if mime_type == 'text/plain':
                # Check file extension for common types that magic misidentifies
                file_ext = Path(file_path).suffix.lower()
                if file_ext == '.csv':
                    mime_type = 'text/csv'
                elif file_ext == '.txt':
                    mime_type = 'text/plain'
                elif file_ext in ['.xlsx', '.xls']:
                    mime_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' if file_ext == '.xlsx' else 'application/vnd.ms-excel'
                elif file_ext == '.docx':
                    mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif file_ext == '.pdf':
                    mime_type = 'application/pdf'
                elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                    mime_type = f'image/{file_ext[1:]}'
            
            return mime_type
        except Exception as e:
            logger.warning(f"⚠️ Magic detection failed: {e}, trying mimetypes")
            # Fallback to mimetypes
            mime_type, _ = mimetypes.guess_type(file_path)
            return mime_type or 'application/octet-stream'
    
    def process_file(self, file_path: str, filename: str = None) -> Dict[str, Any]:
        """
        Process a file and extract text content.
        
        Args:
            file_path: Path to the file to process
            filename: Original filename (optional)
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            # Detect file type
            mime_type = self.detect_file_type(file_path)
            
            # Get processor function
            processor = self.supported_types.get(mime_type)
            if not processor:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {mime_type}',
                    'mime_type': mime_type,
                    'filename': filename or os.path.basename(file_path)
                }
            
            # Process the file
            logger.info(f"📄 Processing {filename or file_path} as {mime_type}")
            result = processor(file_path)
            
            # Add metadata
            result.update({
                'filename': filename or os.path.basename(file_path),
                'mime_type': mime_type,
                'file_size': os.path.getsize(file_path)
            })
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error processing file {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'filename': filename or os.path.basename(file_path)
            }
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using pdfplumber with OCR fallback."""
        try:
            text_content = []
            tables = []
            has_text = False
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content.append(f"--- Page {page_num} ---\n{page_text}")
                        has_text = True
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        for table_num, table in enumerate(page_tables, 1):
                            tables.append({
                                'page': page_num,
                                'table': table_num,
                                'data': table
                            })
            
            # If no text was extracted, try OCR
            if not has_text:
                logger.info("📷 No text found, attempting OCR...")
                ocr_text = self._ocr_pdf(file_path)
                if ocr_text:
                    text_content = [ocr_text]
                    has_text = True
            
            return {
                'success': True,
                'text': '\n\n'.join(text_content),
                'tables': tables,
                'page_count': len(pdf.pages),
                'extraction_method': 'text' if has_text else 'none',
                'has_tables': len(tables) > 0
            }
            
        except Exception as e:
            logger.error(f"❌ PDF processing failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def _ocr_pdf(self, file_path: str) -> str:
        """Perform OCR on PDF pages."""
        try:
            self._init_ocr()
            if not self.ocr_reader:
                return ""
            
            # Convert PDF to images and OCR each page
            import fitz  # PyMuPDF for PDF to image conversion
            
            doc = fitz.open(file_path)
            ocr_results = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Convert to PIL Image
                img = Image.open(io.BytesIO(img_data))
                
                # Perform OCR
                results = self.ocr_reader.readtext(img_data)
                page_text = ' '.join([result[1] for result in results])
                
                if page_text.strip():
                    ocr_results.append(f"--- Page {page_num + 1} (OCR) ---\n{page_text}")
            
            doc.close()
            return '\n\n'.join(ocr_results)
            
        except Exception as e:
            logger.error(f"❌ OCR failed: {e}")
            return ""
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Word document."""
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables = []
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    'table': table_num,
                    'data': table_data
                })
            
            return {
                'success': True,
                'text': '\n\n'.join(paragraphs),
                'tables': tables,
                'paragraph_count': len(paragraphs),
                'table_count': len(tables)
            }
            
        except Exception as e:
            logger.error(f"❌ DOCX processing failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Extract data from Excel file (.xlsx)."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to text representation
                sheet_text = f"=== Sheet: {sheet_name} ===\n"
                sheet_text += f"Columns: {', '.join(df.columns.astype(str))}\n"
                sheet_text += f"Rows: {len(df)}\n\n"
                
                # Add sample data (first 10 rows)
                if not df.empty:
                    sheet_text += "Sample Data:\n"
                    sheet_text += df.head(10).to_string(index=False)
                
                sheets_data[sheet_name] = {
                    'text': sheet_text,
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist()
                }
            
            # Combine all sheets
            all_text = '\n\n'.join([data['text'] for data in sheets_data.values()])
            
            return {
                'success': True,
                'text': all_text,
                'sheets': sheets_data,
                'sheet_count': len(excel_file.sheet_names)
            }
            
        except Exception as e:
            logger.error(f"❌ Excel processing failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def _process_excel_old(self, file_path: str) -> Dict[str, Any]:
        """Extract data from old Excel file (.xls)."""
        try:
            # Read with xlrd engine for .xls files
            excel_file = pd.ExcelFile(file_path, engine='xlrd')
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name, engine='xlrd')
                
                sheet_text = f"=== Sheet: {sheet_name} ===\n"
                sheet_text += f"Columns: {', '.join(df.columns.astype(str))}\n"
                sheet_text += f"Rows: {len(df)}\n\n"
                
                if not df.empty:
                    sheet_text += "Sample Data:\n"
                    sheet_text += df.head(10).to_string(index=False)
                
                sheets_data[sheet_name] = {
                    'text': sheet_text,
                    'rows': len(df),
                    'columns': len(df.columns)
                }
            
            all_text = '\n\n'.join([data['text'] for data in sheets_data.values()])
            
            return {
                'success': True,
                'text': all_text,
                'sheets': sheets_data,
                'sheet_count': len(excel_file.sheet_names)
            }
            
        except Exception as e:
            logger.error(f"❌ XLS processing failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract data from CSV file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                return {'success': False, 'error': 'Could not decode CSV file'}
            
            # Create text representation
            text = f"CSV File Summary:\n"
            text += f"Columns: {', '.join(df.columns.astype(str))}\n"
            text += f"Rows: {len(df)}\n\n"
            text += "Sample Data (first 10 rows):\n"
            text += df.head(10).to_string(index=False)
            
            return {
                'success': True,
                'text': text,
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist()
            }
            
        except Exception as e:
            logger.error(f"❌ CSV processing failed: {e}")
            return {'success': False, 'error': str(e)}
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image using OCR."""
        try:
            self._init_ocr()
            if not self.ocr_reader:
                return {'success': False, 'error': 'OCR not available'}
            
            # Perform OCR
            results = self.ocr_reader.readtext(file_path)
            
            # Extract text with confidence scores
            text_blocks = []
            for (bbox, text, confidence) in results:
                if confidence > 0.5:  # Filter low-confidence results
                    text_blocks.append(text)
            
            extracted_text = ' '.join(text_blocks)
            
            return {
                'success': True,
                'text': extracted_text,
                'ocr_blocks': len(results),
                'high_confidence_blocks': len(text_blocks),
                'extraction_method': 'ocr'
            }
            
        except Exception as e:
            logger.error(f"❌ Image OCR failed: {e}")
            return {'success': False, 'error': str(e)}

# Global instance
document_processor = DocumentProcessor()
